import os
import sys
import re
import time
import uuid
import random
import threading
import webbrowser
import json
import tkinter as tk
from tkinter import scrolledtext, messagebox, ttk, simpledialog, StringVar
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from urllib.parse import urlparse
from datetime import datetime

# ------------------- 路径处理（兼容打包和开发） -------------------
if getattr(sys, 'frozen', False):
    SCRIPT_DIR = os.path.dirname(os.path.abspath(sys.executable))
    RESOURCE_DIR = sys._MEIPASS
else:
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    RESOURCE_DIR = SCRIPT_DIR

# ------------------- 配置管理 -------------------
CONFIG_FILE = os.path.join(SCRIPT_DIR, "config.json")
DEFAULT_CONFIG = {
    "text_patterns": [
        "点击上方 “荆门招商” 可以订阅哦！",
        "（本文部分图片、文字素材来源于网络，"
    ],
    "image_patterns": [
        "/mmbiz_jpg/HyUF4xuREy9OSjBeRmJINQQ6FGmkm7JCqLClJm2OfTpQ9Iq3BzWy4sSandyYQFsyvzx9DggbLlxL8gQ5MvqJrzLe0lowZqLnofqNqVZWvVQ/640",
        "/sz_mmbiz_png/HyUF4xuREyibomAUUYxclO9p5sREOgyLCthJyicYLGoOquibwlf6KRpnELUkBjLH2XpJCblibT10RRsKkOEKRWMDlt3nm4yYibsN7EQIyuhic0TfI/640",
        "/mmbiz_png/bL2iaicTYdZn7Z9T1w7QibEKPdLEZtfxOS5DzUR8icHxXsxiciaNVEkrE8MpkegocurwVAEibkNHk28mCicLGluSbwkiajg/640"
    ]
}

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return DEFAULT_CONFIG.copy()

def save_config(config):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

# 全局配置
CONFIG = load_config()

# ------------------- 全局变量（用于设置窗口交互） -------------------
text_patterns = CONFIG.get("text_patterns", [])
image_patterns = CONFIG.get("image_patterns", [])

# ------------------- 文本清洗函数（使用配置） -------------------
def clean_article_text(raw_text):
    text = raw_text.strip()
    # 删除配置中的文本（按字面匹配，自动转义正则）
    for pattern in text_patterns:
        # 使用 re.escape 转义，按字面匹配，防止用户输入特殊字符导致正则错误
        text = re.sub(re.escape(pattern), '', text, flags=re.DOTALL)
    # 原有删除多余空格和换行逻辑
    text = re.sub(r'[ \t\u3000]{2,}', '\n\n', text)
    paragraphs = re.split(r'\n{2,}', text)
    cleaned = []
    for para in paragraphs:
        para = re.sub(r'\n+', ' ', para)
        para = re.sub(r'[ \t\u3000]+', '', para)
        para = para.strip()
        if para:
            cleaned.append(para)
    return cleaned

# ------------------- 图片过滤（使用配置） -------------------
def is_blocked_image(img_url):
    for pattern in image_patterns:
        if pattern in img_url:
            return True
    return False

# ------------------- 以下为原有函数（未变） -------------------
def get_random_ua():
    USER_AGENTS = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    ]
    return random.choice(USER_AGENTS)

def download_img(img_url):
    if is_blocked_image(img_url):
        print(f"  跳过无用图片: {img_url[:80]}...")
        return None
    try:
        headers = {
            "User-Agent": get_random_ua(),
            "Referer": "https://mp.weixin.qq.com/"
        }
        resp = requests.get(img_url, headers=headers, timeout=10, stream=True)
        resp.raise_for_status()
        content_type = resp.headers.get('content-type', '')
        ext = None
        if 'image/jpeg' in content_type:
            ext = '.jpg'
        elif 'image/png' in content_type:
            ext = '.png'
        elif 'image/gif' in content_type:
            ext = '.gif'
        elif 'image/webp' in content_type:
            ext = '.webp'
        else:
            fmt_match = re.search(r'wx_fmt=([a-zA-Z]+)', img_url)
            if fmt_match:
                fmt = fmt_match.group(1).lower()
                if fmt in ('jpeg', 'jpg'):
                    ext = '.jpg'
                elif fmt == 'png':
                    ext = '.png'
                elif fmt == 'gif':
                    ext = '.gif'
                elif fmt == 'webp':
                    ext = '.webp'
        if not ext:
            ext = '.jpg'
        unique_id = f"{uuid.uuid4().hex}_{int(time.time() * 1000)}"
        filename = f"{unique_id}{ext}"
        path = os.path.join(IMAGE_DIR, filename)
        with open(path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)
        return path
    except Exception as e:
        print(f"图片下载失败 {img_url}, {e}")
        return None

def parse_wx_article(article_url, max_retries=3):
    for attempt in range(1, max_retries + 1):
        try:
            headers = {
                "User-Agent": get_random_ua(),
                "Referer": "https://mp.weixin.qq.com/"
            }
            print(f"  尝试第 {attempt} 次请求...")
            resp = requests.get(article_url, headers=headers, timeout=10)
            resp.encoding = "utf-8"
            soup = BeautifulSoup(resp.text, "lxml")
            title_tag = soup.find("h1", class_="rich_media_title")
            title = title_tag.get_text(strip=True) if title_tag else "无标题"
            content_div = soup.find("div", class_="rich_media_content")
            if not content_div:
                print(f"  第 {attempt} 次：未找到内容区域")
                if attempt < max_retries:
                    time.sleep(2)
                    continue
                else:
                    return None
            raw_text = content_div.get_text("\n", strip=False)
            clean_paragraphs = clean_article_text(raw_text)
            img_paths = []
            for img_tag in content_div.find_all("img"):
                src = img_tag.get("data-src") or img_tag.get("src")
                if src and src.startswith("http"):
                    local_path = download_img(src)
                    if local_path:
                        img_paths.append(local_path)
            return {"title": title, "paragraphs": clean_paragraphs, "images": img_paths}
        except Exception as e:
            print(f"  第 {attempt} 次请求异常：{e}")
            if attempt < max_retries:
                time.sleep(2)
            else:
                return None
    return None

def apply_body_style(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(32)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(16)

def merge_to_word(articles_data, img_width_cm=8):
    today = datetime.now().strftime("%Y%m%d")
    output_file = os.path.join(SCRIPT_DIR, f"微信公众号爬取_{today}.docx")
    doc = Document()
    for data in articles_data:
        p = doc.add_paragraph()
        run = p.add_run(f"标题：{data['title']}")
        run.bold = True
        run.font.size = Pt(18)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for para_text in data['paragraphs']:
            p = doc.add_paragraph(para_text)
            apply_body_style(p)
        for img_path in data['images']:
            try:
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(img_path, width=Cm(img_width_cm))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                apply_body_style(p)
            except Exception as e:
                print(f"插入图片失败 {img_path}: {e}")
    doc.save(output_file)
    print(f"✅ 所有文章已合并保存为 {output_file}")

# ------------------- 设置对话框 -------------------
class SettingsWindow:
    def __init__(self, parent):
        self.parent = parent
        self.window = tk.Toplevel(parent)
        self.window.title("删除规则设置")
        self.window.geometry("600x500")
        self.window.transient(parent)
        self.window.grab_set()

        # 文本规则
        tk.Label(self.window, text="要删除的文本片段（每行一条，支持多行输入）", font=("微软雅黑", 10)).pack(pady=5)
        self.text_listbox = tk.Listbox(self.window, height=6, selectmode=tk.SINGLE)
        self.text_listbox.pack(padx=10, pady=5, fill=tk.X)
        self.load_text_list()
        text_btn_frame = tk.Frame(self.window)
        text_btn_frame.pack(pady=5)
        tk.Button(text_btn_frame, text="添加文本", command=self.add_text_pattern).pack(side=tk.LEFT, padx=5)
        tk.Button(text_btn_frame, text="删除选中", command=self.del_text_pattern).pack(side=tk.LEFT, padx=5)

        # 图片规则
        tk.Label(self.window, text="要过滤的图片URL特征片段（每行一条）", font=("微软雅黑", 10)).pack(pady=5)
        self.image_listbox = tk.Listbox(self.window, height=6, selectmode=tk.SINGLE)
        self.image_listbox.pack(padx=10, pady=5, fill=tk.X)
        self.load_image_list()
        img_btn_frame = tk.Frame(self.window)
        img_btn_frame.pack(pady=5)
        tk.Button(img_btn_frame, text="添加图片特征", command=self.add_image_pattern).pack(side=tk.LEFT, padx=5)
        tk.Button(img_btn_frame, text="删除选中", command=self.del_image_pattern).pack(side=tk.LEFT, padx=5)

        # 保存按钮
        tk.Button(self.window, text="保存配置", command=self.save_settings, bg="#4CAF50", fg="white").pack(pady=15)

    def load_text_list(self):
        self.text_listbox.delete(0, tk.END)
        for item in text_patterns:
            self.text_listbox.insert(tk.END, item)

    def load_image_list(self):
        self.image_listbox.delete(0, tk.END)
        for item in image_patterns:
            self.image_listbox.insert(tk.END, item)

    def add_text_pattern(self):
        # 使用简易多行输入对话框
        dialog = tk.Toplevel(self.window)
        dialog.title("添加要删除的文本")
        dialog.geometry("400x200")
        tk.Label(dialog, text="输入要删除的文本（可直接粘贴）：").pack(pady=5)
        text_widget = scrolledtext.ScrolledText(dialog, height=5)
        text_widget.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        def confirm():
            val = text_widget.get(1.0, tk.END).strip()
            if val:
                text_patterns.append(val)
                self.load_text_list()
            dialog.destroy()
        tk.Button(dialog, text="确认", command=confirm).pack(pady=5)

    def del_text_pattern(self):
        sel = self.text_listbox.curselection()
        if sel:
            index = sel[0]
            del text_patterns[index]
            self.load_text_list()

    def add_image_pattern(self):
        val = simpledialog.askstring("添加图片特征", "输入图片URL中包含的特征片段：", parent=self.window)
        if val:
            image_patterns.append(val)
            self.load_image_list()

    def del_image_pattern(self):
        sel = self.image_listbox.curselection()
        if sel:
            index = sel[0]
            del image_patterns[index]
            self.load_image_list()

    def save_settings(self):
        # 更新全局配置并保存
        CONFIG["text_patterns"] = text_patterns
        CONFIG["image_patterns"] = image_patterns
        save_config(CONFIG)
        messagebox.showinfo("设置", "配置已保存，下次爬取将生效。")
        self.window.destroy()

# ------------------- 主 GUI 应用程序 -------------------
class WxSpiderApp:
    def __init__(self, root):
        self.root = root
        root.title("微信公众号文章批量爬取工具")
        root.geometry("720x650")
        root.resizable(True, True)

        tk.Label(root, text="微信公众号文章批量爬取工具", font=("微软雅黑", 16)).pack(pady=10)
        tk.Label(root, text="请在手机微信中打开公众号 -> 全部消息 -> 复制文章链接（每行一个）",
                 fg="blue", font=("微软雅黑", 10)).pack()

        self.text_area = scrolledtext.ScrolledText(root, height=12, font=("微软雅黑", 10))
        self.text_area.pack(padx=20, pady=10, fill=tk.BOTH, expand=True)

        # 图片宽度设置
        width_frame = tk.Frame(root)
        width_frame.pack(pady=5)
        tk.Label(width_frame, text="图片宽度 (cm)：", font=("微软雅黑", 10)).pack(side=tk.LEFT, padx=5)
        self.width_var = StringVar()
        self.width_var.set("8.0")
        self.width_spinbox = tk.Spinbox(width_frame, from_=2.0, to=15.0, increment=0.5,
                                        textvariable=self.width_var, width=8,
                                        font=("微软雅黑", 10))
        self.width_spinbox.pack(side=tk.LEFT, padx=5)
        tk.Label(width_frame, text="(范围 2.0 ~ 15.0 cm)", font=("微软雅黑", 9), fg="gray").pack(side=tk.LEFT)

        # 按钮框架
        btn_frame = tk.Frame(root)
        btn_frame.pack(pady=10)

        self.start_btn = tk.Button(btn_frame, text="开始爬取", command=self.start_crawl,
                                   width=12, bg="#4CAF50", fg="white", font=("微软雅黑", 10))
        self.start_btn.pack(side=tk.LEFT, padx=5)

        self.help_btn = tk.Button(btn_frame, text="使用帮助", command=self.show_help,
                                  width=12, bg="#2196F3", fg="white", font=("微软雅黑", 10))
        self.help_btn.pack(side=tk.LEFT, padx=5)

        self.settings_btn = tk.Button(btn_frame, text="设置规则", command=self.open_settings,
                                      width=12, bg="#FF9800", fg="white", font=("微软雅黑", 10))
        self.settings_btn.pack(side=tk.LEFT, padx=5)

        self.clear_btn = tk.Button(btn_frame, text="清空链接", command=self.clear_text,
                                   width=12, bg="#f44336", fg="white", font=("微软雅黑", 10))
        self.clear_btn.pack(side=tk.LEFT, padx=5)

        # 进度条
        self.progress = ttk.Progressbar(root, orient=tk.HORIZONTAL, length=400, mode='determinate')
        self.progress.pack(pady=10)

        # 状态栏
        self.status_var = StringVar()
        self.status_var.set("就绪")
        status_label = tk.Label(root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_label.pack(side=tk.BOTTOM, fill=tk.X)

    def clear_text(self):
        self.text_area.delete(1.0, tk.END)

    def show_help(self):
        help_file = os.path.join(RESOURCE_DIR, "help.html")
        if os.path.exists(help_file):
            webbrowser.open(help_file)
        else:
            messagebox.showinfo("帮助", "帮助文件 help.html 未找到，请确保它与程序在同一目录。")

    def open_settings(self):
        SettingsWindow(self.root)

    def start_crawl(self):
        content = self.text_area.get(1.0, tk.END).strip()
        urls = [line.strip() for line in content.split('\n') if line.strip()]
        if not urls:
            messagebox.showwarning("提示", "请至少输入一个有效的文章链接！")
            return
        try:
            img_width = float(self.width_var.get().strip())
            if img_width < 2.0 or img_width > 15.0:
                raise ValueError
        except:
            messagebox.showerror("输入错误", "图片宽度请输入 2.0 ~ 15.0 之间的数值（单位：cm）")
            return

        self.start_btn.config(state=tk.DISABLED)
        self.status_var.set("正在爬取，请稍候...")
        self.progress['maximum'] = len(urls)
        self.progress['value'] = 0
        threading.Thread(target=self.crawl_thread, args=(urls, img_width), daemon=True).start()

    def crawl_thread(self, urls, img_width):
        all_data = []
        total = len(urls)
        try:
            for idx, url in enumerate(urls, 1):
                self.root.after(0, lambda i=idx, t=total: self.status_var.set(f"正在处理第 {i}/{t} 篇..."))
                data = parse_wx_article(url)
                if data:
                    all_data.append(data)
                if idx < total:
                    time.sleep(random.uniform(1, 3))
                self.root.after(0, lambda v=idx: self.progress.config(value=v))

            if all_data:
                self.root.after(0, lambda: self.status_var.set("正在生成Word文档..."))
                merge_to_word(all_data, img_width_cm=img_width)
                self.root.after(0, lambda: messagebox.showinfo("完成",
                             f"成功爬取 {len(all_data)} 篇文章，已生成 Word 文档。\n文件名：微信公众号爬取_{datetime.now().strftime('%Y%m%d')}.docx"))
            else:
                self.root.after(0, lambda: messagebox.showerror("错误", "所有文章解析失败，请检查链接是否有效。"))
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("异常", f"发生错误：{str(e)}"))
        finally:
            self.root.after(0, lambda: self.start_btn.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.status_var.set("就绪"))
            self.root.after(0, lambda: self.progress.config(value=0))

# ------------------- 初始化配置（确保存在） -------------------
if not os.path.exists(CONFIG_FILE):
    save_config(DEFAULT_CONFIG)

# ------------------- 启动 -------------------
if __name__ == "__main__":
    root = tk.Tk()
    app = WxSpiderApp(root)
    root.mainloop()